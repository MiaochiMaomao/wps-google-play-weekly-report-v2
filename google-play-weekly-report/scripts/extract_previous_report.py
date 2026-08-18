from __future__ import annotations

import argparse
import json
from pathlib import Path
from zipfile import ZipFile

from docx import Document


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    source = Path(args.input).resolve()
    output = Path(args.output).resolve()
    doc = Document(source)
    result = {
        "source": str(source),
        "paragraphs": [
            {"index": i, "style": p.style.name if p.style else "", "text": p.text}
            for i, p in enumerate(doc.paragraphs)
            if p.text.strip()
        ],
        "tables": [
            {
                "index": ti,
                "rows": [[cell.text for cell in row.cells] for row in table.rows],
            }
            for ti, table in enumerate(doc.tables)
        ],
        "sections": [
            {
                "index": i,
                "page_width": int(s.page_width),
                "page_height": int(s.page_height),
                "left_margin": int(s.left_margin),
                "right_margin": int(s.right_margin),
                "top_margin": int(s.top_margin),
                "bottom_margin": int(s.bottom_margin),
            }
            for i, s in enumerate(doc.sections)
        ],
        "inline_shapes": len(doc.inline_shapes),
    }
    with ZipFile(source) as zf:
        result["media"] = [
            {"name": name, "size": zf.getinfo(name).file_size}
            for name in zf.namelist()
            if name.startswith("word/media/")
        ]
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({
        "paragraphs": len(result["paragraphs"]),
        "tables": len(result["tables"]),
        "sections": len(result["sections"]),
        "inline_shapes": result["inline_shapes"],
        "media": len(result["media"]),
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()
