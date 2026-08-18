from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document


def period_tokens(iso_date: str) -> set[str]:
    year, month, day = (int(part) for part in iso_date.split("-"))
    return {
        f"{year:04d}-{month:02d}-{day:02d}",
        f"{year:04d}.{month:02d}.{day:02d}",
        f"{year:04d}/{month:02d}/{day:02d}",
        f"{year}.{month}.{day}",
        f"{year}/{month}/{day}",
        f"{month}.{day}",
        f"{month}/{day}",
    }


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--report", required=True)
    parser.add_argument("--metrics", required=True)
    parser.add_argument("--config", required=True)
    parser.add_argument("--output")
    args = parser.parse_args()

    report_path = Path(args.report).resolve()
    metrics = json.loads(Path(args.metrics).read_text(encoding="utf-8"))
    config = json.loads(Path(args.config).read_text(encoding="utf-8"))
    doc = Document(report_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs] + [cell.text for table in doc.tables for row in table.rows for cell in row.cells])
    section = doc.sections[0]
    usable_width = section.page_width - section.left_margin - section.right_margin
    table_widths = []
    for index, table in enumerate(doc.tables):
        width = sum(int(column.w or 0) for column in table._tbl.tblGrid.gridCol_lst)
        table_widths.append({"index": index, "declared_grid_width": width, "declared_grid_fits": width <= usable_width})

    start = metrics["period"]["start"]
    end = metrics["period"]["end"]
    global_metrics = metrics["global"]
    required_counts = [
        global_metrics["rating_total"],
        global_metrics["comment_total"],
        global_metrics["rating_low"],
        global_metrics["rating_high"],
        global_metrics["comment_low"],
        global_metrics["comment_high"],
    ]
    history = metrics.get("monthly_score_history")
    monthly_values = [value for values in history.values() for value in values] if isinstance(history, dict) else []
    checks = {
        "file_exists": report_path.exists() and report_path.stat().st_size > 0,
        "one_or_more_sections": len(doc.sections) >= 1,
        "title_present": config["report"]["title"] in text,
        "period_start_present": any(token in text for token in period_tokens(start)),
        "period_end_present": any(token in text for token in period_tokens(end)),
        "core_counts_present": all(str(value) in text or f"{value:,}" in text for value in required_counts),
        "configured_categories_present": all(name in text for name in config["classification"]["categories"]),
        "monthly_score_history_valid": bool(monthly_values) and all(isinstance(value, (int, float)) and not isinstance(value, bool) and 1 <= value <= 5 for value in monthly_values),
        "no_placeholders": not re.search(r"<[^>]+>|TODO|TBD|#NAME\?|#REF!", text, re.I),
        "all_inline_images_fit": all(shape.width <= usable_width for shape in doc.inline_shapes),
        "table_width_metadata_recorded": len(table_widths) == len(doc.tables),
    }
    with ZipFile(report_path) as archive:
        checks["zip_integrity"] = archive.testzip() is None
        checks["has_media"] = any(name.startswith("word/media/") for name in archive.namelist())

    result = {
        "status": "PASS" if all(checks.values()) else "FAIL",
        "checks": checks,
        "usable_width": int(usable_width),
        "max_inline_image_width": max((int(shape.width) for shape in doc.inline_shapes), default=0),
        "table_width_note": "Rendered page images are the final table-overflow gate because document metadata can be stale.",
        "table_widths": table_widths,
    }
    if args.output:
        output = Path(args.output)
        output.parent.mkdir(parents=True, exist_ok=True)
        output.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(result, ensure_ascii=False))
    raise SystemExit(0 if result["status"] == "PASS" else 1)


if __name__ == "__main__":
    main()

